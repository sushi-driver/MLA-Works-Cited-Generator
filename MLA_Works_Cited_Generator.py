import os
import docx
import tkinter
from tkinter import *
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import Length
from titlecase import titlecase
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

root = Tk()
root.title("iFormat")
doc = docx.Document("MLA_Works_Cited.docx")

def clickgen1(): # Generate 'Print Book' cite and add to document
    author = authorl1.get().title() + ", " + authorf1.get().title() + ". "
    newtitle = titlecase(title1.get()) + ". "
    neweditor = "Edited by " + editor1.get().title() + ", "
    newpublisher = titlecase(publisher1.get()) + ", "
    newpubyear = pubyear1.get() + "."

    if len(author) < 5:
        author = ""
    if len(newtitle) < 3:
        newtitle = ""
    if len(neweditor) < 13:
        neweditor = ""
    if len(newpublisher) < 4:
        newpublisher = ""
    if len(newpubyear) < 5:
        newpubyear = ""
           
    p = doc.add_paragraph(author)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.style = "MLAWC"
    p.add_run(newtitle).italic = True
    p.add_run(neweditor + newpublisher + newpubyear)
    doc.save("MLA_Works_Cited.docx")
    os.system('start MLA_Works_Cited.docx')

def clickgen1_1(): # Generate 2 author 'Print Book' cite and add to document 
    author = authorl1_1.get().title() + ", " + authorf1_1.get().title() + ", "
    secauth = "and " + titlecase(secauth_1.get()) + ". "
    newtitle = titlecase(title1_1.get()) + ". "
    neweditor = "Edited by " + editor1_1.get().title() + ", "
    newpublisher = titlecase(publisher1_1.get()) + ", "
    newpubyear = pubyear1_1.get() + "."

    if len(author) < 5:
        author = ""
    if len(secauth) < 7:
        secauth = ""
    if len(newtitle) < 3:
        newtitle = ""
    if len(neweditor) < 13:
        neweditor = ""
    if len(newpublisher) < 4:
        newpublisher = ""
    if len(newpubyear) < 5:
        newpubyear = ""
           
    p = doc.add_paragraph(author + secauth)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.style = "MLAWC"
    p.add_run(newtitle).italic = True
    p.add_run(neweditor + newpublisher + newpubyear)
    doc.save("MLA_Works_Cited.docx")
    os.system('start MLA_Works_Cited.docx')

def clickgen1_2(): # Generate 3 or more author 'Print Book' cite and add to document
    author = authorl1_2.get().title() + ", " + authorf1_2.get().title() + ", et al. "
    newtitle = titlecase(title1_2.get()) + ". "
    neweditor = "Edited by " + editor1_2.get().title() + ", "
    newpublisher = titlecase(publisher1_2.get()) + ", "
    newpubyear = pubyear1_2.get() + "."

    if len(author) < 12:
        author = ""
    if len(newtitle) < 3:
        newtitle = ""
    if len(neweditor) < 13:
        neweditor = ""
    if len(newpublisher) < 4:
        newpublisher = ""
    if len(newpubyear) < 5:
        newpubyear = ""
           
    p = doc.add_paragraph(author)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.style = "MLAWC"
    p.add_run(newtitle).italic = True
    p.add_run(neweditor + newpublisher + newpubyear)
    doc.save("MLA_Works_Cited.docx")
    os.system('start MLA_Works_Cited.docx')

def clickgen2(): # Generate 'Print Article' cite and add to document
    author = authorl2.get().title() + ", " + authorf2.get().title() + "."
    newart_title = " \"" + titlecase(art_title2.get()) + ".\" "
    newper_title = titlecase(per_title2.get()) 
    newvolume = "vol. " + volume2.get() + ", "
    newnumber = "no. " + number2.get() + ", "
    month = (month2.get()[slice(0, 3)])
    date = day2.get() + ", " + month.title() + ". " + year2.get() + ", "
    pages = "pp. " + pgstart2.get() + "-" + pgend2.get() + ". "
    if len(author) < 5:
        author = ""
    if len(newart_title) < 6:
        newart_title = ""
    if len(newper_title) < 1:
        newper_title = ""
    if len(newvolume) < 8:
        newvolume = ""
    if len(newnumber) < 7:
        newnumber = ""
    if len(date) < 14:
        date = date.replace(", ", "", 1)
    if len(date) < 12:
        date = year2.get() + ", "
    if len(date) < 5:
        date = ""
    if len(pages) < 8:
        pages = ""

    p = doc.add_paragraph(author + newart_title)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.style = "MLAWC"
    p.add_run(newper_title).italic = True
    if len(newper_title) > 7:
        p.add_run(", ")
    p.add_run(newvolume + newnumber + date + pages)
    doc.save("MLA_Works_Cited.docx")
    os.system('start MLA_Works_Cited.docx')
    
def clickgen3(): # Generate 'Website' cite and add to document
    newweb_title = title3.get().capitalize() + ". "
    newwebpub = titlecase(publisher3.get()) + ", "
    date = year3.get() + "-" + year3_1.get() + ", "
    newurl = str(url3.get()) + "."

    if len(newweb_title) < 3:
        newweb_title = ""
    if len(newwebpub) < 3:
        newwebpub = ""
    if len(date) < 8:
        date = date.lstrip("-")
    if len(date) < 6:
        date = ""
    if len(newurl) < 4:
        newurl = ""

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.style = "MLAWC"
    p.add_run(newweb_title).italic = True
    p.add_run(newwebpub + date + newurl)
    doc.save("MLA_Works_Cited.docx")
    os.system('start MLA_Works_Cited.docx')   

def clickgen4(): # Generate 'Web Article' cite and add to document
    author = authorl4.get().title() + ", " + authorf4.get().title() + ". "
    newart_title = "\"" + titlecase(title4.get()) + ".\" "
    newwebsite = titlecase(website4.get())
    pubyear = pubyear4.get() + ", "
    month = (month4.get()[slice(0, 3)])
    newurl = str(url4.get()) + ". "
    date = day4.get() + ", " + month.title() + ". " + year4.get() + "."
    
    if len(author) < 5:
        author = ""
    if len(newart_title) < 5:
        newart_title = ""
    if len(newwebsite) < 3:
        newwebpub = ""
    if len(pubyear) < 4:
        pubyear = ""
    if len(newurl) < 4:
        newurl = ""
    if len(date) < 13:
        date = date.replace(", ", "", 1)
    if len(date) < 10:
        date = year4.get() + "."
    if len(date) < 5:
        date = ""
    
    p = doc.add_paragraph(author + newart_title)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.style = "MLAWC"
    p.add_run(newwebsite).italic = True
    if len(newwebsite) > 3:
        p.add_run(", ")
    p.add_run(pubyear + newurl + date)
    doc.save("MLA_Works_Cited.docx")
    os.system('start MLA_Works_Cited.docx')  

def clickgen5(): # Generate 'Online Video' cite and add to doc
    author = authorl5.get().title() + ", " + authorf5.get().title() + ", "
    vid = "\"" + titlecase(vid5.get()) + ".\" " 
    website = website5.get()
    month = (month5.get()[slice(0, 3)])
    pubdate = day5.get() + ", " + month.title() + ". " + year5.get() + ", "
    url = url5.get() + "."

    if len(author) < 5:
        author = ""
    if len(vid) < 5:
        vid = ""
    if len(website) < 1:
        website = ""
    if len(month) < 2:
        month = ""
    if len(pubdate) < 14:
        pubdate = pubdate.replace(", ", "", 1)
    if len(pubdate) < 11:
        pubdate = year5.get() + ", "
    if len(pubdate) < 6:
        pubdate = ""
    if len(url) < 3:
        url = ""

    p = doc.add_paragraph(author + vid)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.style = "MLAWC"
    p.add_run(website).italic = True
    if len(website) > 3:
        p.add_run(", ")
    p.add_run(pubdate + url)
    doc.save("MLA_Works_Cited.docx")
    os.system('start MLA_Works_Cited.docx')  
    
def clickgen6(): # Generate 'Online Database' cite and add to doc
    author = authorl6.get().title() + ", " + authorf6.get().title() + ". "
    newwork = "\"" + titlecase(work6.get()) + ".\" "
    newper_title = titlecase(per_title6.get())
    newvlm = "vol. " + volume6.get() + ", "
    newiss = "no. " + issue6.get() +", "
    newyear = year6.get() + ", "
    pages = "pp. " + pgstart6.get() + "-" + pgend6.get() + ". "
    title = titlecase(title6.get())
    newdoi = doi6.get() + "."

    if len(author) < 5:
        author = ""
    if len(newwork) < 5:
        newwork = ""
    if len(newper_title) < 1:
        newper_title = ""
    if len(newvlm) < 8:
        newvlm = ""
    if len(newiss) < 7:
        newiss = ""
    if len(newyear) < 5:
        newyear = ""
    if len(pages) < 8:
        pages  = ""
    if len(title) < 1:
        title = ""
    if len(newdoi) < 3:
        newdoi = ""

    p = doc.add_paragraph(author + newwork)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.style = "MLAWC"
    p.add_run(newper_title).italic = True
    if len(newper_title) > 2:
        p.add_run(", ")
    p.add_run(newvlm + newiss + newyear)
    p.add_run(title).italic = True
    if len(title) > 2:
        p.add_run(", ")
    p.add_run(newdoi)
    doc.save("MLA_Works_Cited.docx")
    os.system('start MLA_Works_Cited.docx')

def clickgen7(): # Generate 'Film or Movie' cite and add to document
    perf = perfl7.get().title() + ", " + perff7.get().title() + ", performer. "
    title = titlecase(title7.get())
    dir = "Directed by " + dir7.get().title() + ", "
    stud = titlecase(stud7.get()) + ", "
    year = year7.get() + "."

    if len(perf) < 16:
        perf = ""
    if len(title) < 1:
        title = ""
    if len(dir) < 15:
        dir = ""
    if len(stud) < 3:
        stud = ""
    if len(year) < 2:
        year = "" 

    p=doc.add_paragraph(perf)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.style = "MLAWC"
    p.add_run(title).italic = True
    if len(title) > 2:
        p.add_run(", ")
    p.add_run(dir + stud + year)
    doc.save("MLA_Works_Cited.docx")
    os.system('start MLA_Works_Cited.docx')

def clickgen8(): # Generate 'Wikipedia' cite and add to document
    sub = "\"" + titlecase(sub8.get()) + ".\" " 
    month = (month8.get()[slice(0, 3)])
    date = day8.get() + ", " + month.title() + ". " + year8.get() + ", "
    url = url8.get() + "."

    if len(sub) < 5:
        sub = ""
    if len(month) < 3:
        month = ""
    if len(date) < 13:
        date = date.replace(", ", "", 1)
    if len(date) < 11:
        date = year8.get() + ", "
    if len(date) < 5:
        date = ""
    if len(url) < 2:
        url = ""

    p = doc.add_paragraph(sub)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.style = "MLAWC"
    p.add_run("Wikipedia, ").italic = True
    p.add_run(date + url)
    doc.save("MLA_Works_Cited.docx")
    os.system('start MLA_Works_Cited.docx')  

def click1(): # Print Book
    # Generate drop down for number of authors
    global num
    global top0 
    top0 = Toplevel()
    top0.title("How Many Authors?")
    x = 0
    count = ["1 ", "2 ", "3"]
    authcount = StringVar(top0)
    authcount.set(count[0])
    label00 = Label(top0, text = "How many authors does your book have?", font="calibri 14 italic")
    label101 = Label(top0, text = "(For 3 or more authors, please select 3)", font="calibri 12 italic") 
    label00.grid(row = 0, columnspan = 2, padx = 5, pady = 5)
    label101.grid(row = 1, columnspan = 2)
    label0 = OptionMenu(top0, authcount, *count)
    label0.configure(width = 25, height = 1)
    label0.grid(row = 2, column = 0)
    def ok(): # Get number of authors
        num = authcount.get()
        x = int(num)
        if x == 1: # One author
            global top1_1
            top1_1 = Toplevel()
            top1_1.title("Print Book...Capitalization and formatting will be automatically applied")

            label1 = Label(top1_1, text = "Author's First Name:", font="calibri 12")
            label1.grid(row = 0)
            global authorf1
            authorf1 = Entry(top1_1, width = 100)
            authorf1.grid(row = 0, column = 1, padx = 2, pady = 2)

            label2 = Label(top1_1, text = "Author's Last Name:", font="calibri 12")
            label2.grid(row = 1)
            global authorl1
            authorl1 = Entry(top1_1, width = 100)
            authorl1.grid(row = 1, column = 1, padx = 2, pady = 2)

            label3 = Label(top1_1, text = "Title and Subtitle:", font="calibri 12")
            label3.grid(row = 2)
            global title1
            title1 = Entry(top1_1, width = 100)
            title1.grid(row = 2, column = 1, padx = 2, pady = 2)

            label4 = Label(top1_1, text = "Publisher:", font="calibri 12")
            label4.grid(row = 3)
            global publisher1
            publisher1 = Entry(top1_1, width = 100)
            publisher1.grid(row = 3, column = 1, padx = 2, pady = 2)

            label5 = Label(top1_1, text = "Year of Publication:", font="calibri 12")
            label5.grid(row = 4)
            global pubyear1
            pubyear1 = Entry(top1_1, width = 100)
            pubyear1.grid(row = 4, column = 1, padx = 2, pady = 2)

            label6 = Label(top1_1, text = "Editor's Full Name:", font="calibri 12")
            label6.grid(row = 5)
            global editor1
            editor1 = Entry(top1_1, width = 100)
            editor1.grid(row = 5, column = 1, padx = 2, pady = 2)

            generate = Button(top1_1, text = "Generate Works Cited Document", font="calibri 12", command = clickgen1)
            generate.grid(row = 7, column = 1, pady = 5)
            proof = Label(top1_1, text = "Please proofread your document...iFormat is not responsible for your bad grades", font="calibri 10 italic")
            proof.grid(row = 8, column = 1)
    
        if x == 2: # Two authors
            global top1_2
            top1_2 = Toplevel()
            top1_2.title("Print Book...Capitalization and formatting will be automatically applied")

            label1 = Label(top1_2, text = "Main Author's First Name:", font="calibri 12")
            label1.grid(row = 0)
            global authorf1_1
            authorf1_1 = Entry(top1_2, width = 100)
            authorf1_1.grid(row = 0, column = 1, padx = 2, pady = 2)

            label2 = Label(top1_2, text = "Main Author's Last Name:", font="calibri 12")
            label2.grid(row = 1)
            global authorl1_1
            authorl1_1 = Entry(top1_2, width = 100)
            authorl1_1.grid(row = 1, column = 1, padx = 2, pady = 2)

            label2_1 = Label(top1_2, text = "Second Author's Full Name:", font="calibri 12")
            label2_1.grid(row = 2)
            global secauth_1
            secauth_1 = Entry(top1_2, width = 100)
            secauth_1.grid(row = 2, column = 1, padx = 2, pady = 2)

            label3 = Label(top1_2, text = "Title and Subtitle:", font="calibri 12")
            label3.grid(row = 3)
            global title1_1
            title1_1 = Entry(top1_2, width = 100)
            title1_1.grid(row = 3, column = 1, padx = 2, pady = 2)

            label4 = Label(top1_2, text = "Publisher:", font="calibri 12")
            label4.grid(row = 4)
            global publisher1_1
            publisher1_1 = Entry(top1_2, width = 100)
            publisher1_1.grid(row = 4, column = 1, padx = 2, pady = 2)

            label5 = Label(top1_2, text = "Year of Publication:", font="calibri 12")
            label5.grid(row = 5)
            global pubyear1_1
            pubyear1_1 = Entry(top1_2, width = 100)
            pubyear1_1.grid(row = 5, column = 1, padx = 2, pady = 2)

            label6 = Label(top1_2, text = "Editor's Full Name:", font="calibri 12")
            label6.grid(row = 6)
            global editor1_1
            editor1_1 = Entry(top1_2, width = 100)
            editor1_1.grid(row = 6, column = 1, padx = 2, pady = 2)

            generate = Button(top1_2, text = "Generate Works Cited Document", font="calibri 12", command = clickgen1_1)
            generate.grid(row = 7, column = 1, pady = 5)
            proof = Label(top1_2, text = "Please proofread your document...iFormat is not responsible for your bad grades", font="calibri 10 italic")
            proof.grid(row = 8, column = 1)

        if x == 3: # Three or more authors
            global top1_3
            top1_3 = Toplevel()
            top1_3.title("Print Book...Capitalization and formatting will be automatically applied")

            label1 = Label(top1_3, text = "Main Author's First Name:", font="calibri 12")
            label1.grid(row = 0)
            global authorf1_2
            authorf1_2 = Entry(top1_3, width = 100)
            authorf1_2.grid(row = 0, column = 1, padx = 2, pady = 2)

            label2 = Label(top1_3, text = "Main Author's Last Name:", font="calibri 12")
            label2.grid(row = 1)
            global authorl1_2
            authorl1_2 = Entry(top1_3, width = 100)
            authorl1_2.grid(row = 1, column = 1, padx = 2, pady = 2)

            message = Label(top1_3, text = "For works composed by 3 or more authors, the expression \'et al\' will represent all additional authors...", font="calibri 10 italic")
            message.grid(row = 2, columnspan = 2)

            label3 = Label(top1_3, text = "Title and Subtitle:", font="calibri 12")
            label3.grid(row = 3)
            global title1_2
            title1_2 = Entry(top1_3, width = 100)
            title1_2.grid(row = 3, column = 1, padx = 2, pady = 2)

            label4 = Label(top1_3, text = "Publisher:", font="calibri 12")
            label4.grid(row = 4)
            global publisher1_2
            publisher1_2 = Entry(top1_3, width = 100)
            publisher1_2.grid(row = 4, column = 1, padx = 2, pady = 2)

            label5 = Label(top1_3, text = "Year of Publication:", font="calibri 12")
            label5.grid(row = 5)
            global pubyear1_2
            pubyear1_2 = Entry(top1_3, width = 100)
            pubyear1_2.grid(row = 5, column = 1, padx = 2, pady = 2)

            label6 = Label(top1_3, text = "Editor's Full Name:", font="calibri 12")
            label6.grid(row = 6)
            global editor1_2
            editor1_2 = Entry(top1_3, width = 100)
            editor1_2.grid(row = 6, column = 1, padx = 2, pady = 2)

            generate = Button(top1_3, text = "Generate Works Cited Document", font="calibri 12", command = clickgen1_2)
            generate.grid(row = 7, column = 1, pady = 5)   
            proof = Label(top1_3, text = "Please proofread your document...iFormat is not responsible for your bad grades", font="calibri 10 italic")
            proof.grid(row = 8, column = 1)
            
    button = Button(top0, text = "OK", command = ok)
    button.grid(row = 2, column = 1, padx = 2, pady = 5)
           
def click2(): # Print Article
    global top2 
    top2 = Toplevel()
    top2.title("Print Article...Capitalization and formatting will be automatically applied")

    label1 = Label(top2, text = "Author's First Name:", font="calibri 12")
    label1.grid(row = 0)
    global authorf2
    authorf2 = Entry(top2, width = 100)
    authorf2.grid(row = 0, column = 1, padx = 2, pady = 2)

    label2 = Label(top2, text = "Author's Last Name:", font="calibri 12")
    label2.grid(row = 1)
    global authorl2
    authorl2 = Entry(top2, width = 100)
    authorl2.grid(row = 1, column = 1, padx = 2, pady = 2)

    label3 = Label(top2, text = "Title of Article:", font="calibri 12")
    label3.grid(row = 2)
    global art_title2
    art_title2 = Entry(top2, width = 100)
    art_title2.grid(row = 2, column = 1, padx = 2, pady = 2)

    label4 = Label(top2, text = "Title of Periodical:", font="calibri 12")
    label4.grid(row = 3)
    global per_title2 
    per_title2 = Entry(top2, width = 100)
    per_title2.grid(row = 3, column = 1, padx = 2, pady = 2)

    label5 = Label(top2, text = "Article Volume:", font="calibri 12")
    label5.grid(row = 4) 
    global volume2 
    volume2 = Entry(top2, width = 100)
    volume2.grid(row= 4, column = 1, padx = 2, pady = 2)

    label6 = Label(top2, text = "Article Number:", font="calibri 12")
    label6.grid(row = 5)
    global number2
    number2 = Entry(top2, width = 100)
    number2.grid(row = 5, column = 1, padx = 2, pady = 2)

    frame = Frame(top2)
    frame.grid(row = 6, column = 1, columnspan = 6, sticky = "w")
    label7 = Label(top2, text = "Article Date:", font="calibri 12")
    label7.grid(row = 6, column = 0, padx = 2, pady = 2)

    label8 = Label(frame, text = "Day:").grid(row = 0, column = 1, sticky = "w")
    global day2
    day2 = Entry(frame, width = 10)
    day2.grid(row = 0, column = 2, padx = 2, pady = 2)

    label9 = Label(frame, text = "Month:").grid(row = 0, column = 3)
    global month2 
    month2 = Entry(frame, width = 10)
    month2.grid(row = 0, column = 4, padx = 2, pady = 2)

    label10 = Label(frame, text = "Year:").grid(row = 0, column = 5)
    global year2 
    year2 = Entry(frame, width = 10)
    year2.grid(row = 0, column = 6, padx = 2, pady = 2)

    frame2 = Frame(top2)
    frame2.grid(row = 7, column = 1, columnspan = 4, sticky = "w")
    label11 = Label(top2, text = "Page Numbers:", font="calibri 12")
    label11.grid(row = 7, column = 0, padx = 2, pady = 2)

    label12 = Label(frame2, text = "Start Page:").grid(row = 0, column = 1)
    global pgstart2
    pgstart2 = Entry(frame2, width = 10)
    pgstart2.grid(row = 0, column = 2, padx = 2, pady = 2)

    label13 = Label(frame2, text = "End Page:").grid(row = 0, column = 3)
    global pgend2
    pgend2 = Entry(frame2, width = 10)
    pgend2.grid(row = 0, column = 4, padx = 2, pady = 2)
    
    generate = Button(top2, text = "Generate Works Cited Document", font="calibri 12", command = clickgen2)
    generate.grid(row = 10, column = 1, pady = 5)
    proof = Label(top2, text = "Please proofread your document...iFormat is not responsible for your bad grades", font="calibri 10 italic")
    proof.grid(row = 11, column = 1)
  
def click3(): # Website
    global top3 
    top3 = Toplevel()
    top3.title("Website...Capitalization and formatting will be automatically applied")

    label1 = Label(top3, text = "Website Name/Title:", font="calibri 12")
    label1.grid(row = 2)
    global title3
    title3 = Entry(top3, width = 100)
    title3.grid(row = 2, column = 1, padx = 2, pady = 2)

    label2 = Label(top3, text = "Website Sponsor/Publisher:", font="calibri 12")
    label2.grid(row = 3)
    global publisher3
    publisher3 = Entry(top3, width = 100)
    publisher3.grid(row= 3, column = 1, padx = 2, pady = 2)

    frame = Frame(top3)
    frame.grid(row = 4, column = 1, columnspan = 6, sticky = "w")
    label3 = Label(top3, text = "Date Span:", font="calibri 12")
    label3.grid(row = 4, column = 0, padx = 2, pady = 2)

    label4 = Label(frame, text = "Year of Publication:").grid(row = 0, column = 1, sticky = "w")
    global year3
    year3 = Entry(frame, width = 10)
    year3.grid(row = 0, column = 2, padx = 2, pady = 2)

    label5 = Label(frame, text = "Year Last Updated:").grid(row = 0, column = 3)
    global year3_1 
    year3_1 = Entry(frame, width = 10)
    year3_1.grid(row = 0, column = 4, padx = 2, pady = 2)

    label6 = Label(top3, text = "Website URL:", font="calibri 12")
    label6.grid(row = 5)
    global url3 
    url3 = Entry(top3, width = 100)
    url3.grid(row = 5, column = 1, padx = 2, pady = 2)

    generate = Button(top3, text = "Generate Works Cited Document", font="calibri 12", command = clickgen3)
    generate.grid(row = 6, column = 1, pady = 5)
    proof = Label(top3, text = "Please proofread your document...iFormat is not responsible for your bad grades", font="calibri 10 italic")
    proof.grid(row = 7, column = 1)

def click4(): # Web Article
    def on_entry_click(event): # controls ghost text in 'Year of Publication' field
        if pubyear4.get() == "If unavailabe, use \'Date Accessed\' below. Do NOT fill in both fields":
           pubyear4.delete(0, "end")
           pubyear4.insert(0, "")
           pubyear4.config(fg = "black")
    def on_focusout(event):
        if pubyear4.get() == "":
            pubyear4.insert(0, "If unavailabe, use \'Date Accessed\' below. Do NOT fill in both fields")
            pubyear4.config(fg = "grey")

    global top4 
    top4 = Toplevel()
    top4.title("Web Article...Capitalization and formatting will be automatically applied")

    label1 = Label(top4, text = "Author's First Name:", font="calibri 12")
    label1.grid(row = 0)
    global authorf4
    authorf4 = Entry(top4, width = 100)
    authorf4.grid(row = 0, column = 1, padx = 2, pady = 2)

    label2 = Label(top4, text = "Author's Last Name:", font="calibri 12")
    label2.grid(row = 1)
    global authorl4
    authorl4 = Entry(top4, width = 100)
    authorl4.grid(row = 1, column = 1, padx = 2, pady = 2)

    label3 = Label(top4, text = "Article Title:", font="calibri 12")
    label3.grid(row = 2)
    global title4
    title4 = Entry(top4, width = 100)
    title4.grid(row = 2, column = 1, padx = 2, pady = 2)

    label4 = Label(top4, text = "Website:", font="calibri 12")
    label4.grid(row = 3)
    global website4
    website4 = Entry(top4, width = 100)
    website4.grid(row= 3, column = 1, padx = 2, pady = 2)

    label5 =Label(top4, text = "Year of Publication:", font="calibri 12")
    label5.grid(row = 4)
    global pubyear4 
    pubyear4 = Entry(top4, width = 100)
    pubyear4.grid(row = 4, column = 1, padx = 2, pady = 2)
    pubyear4.insert(0, "If unavailabe, use \'Date Accessed\' below. Do NOT fill in both fields")
    pubyear4.bind("<FocusIn>", on_entry_click)
    pubyear4.bind("<FocusOut>", on_focusout)
    pubyear4.config(fg = "grey")

    frame = Frame(top4)
    frame.grid(row = 5, column = 1, columnspan = 6, sticky = "w")
    label6 = Label(top4, text = "Date Accessed:", font="calibri 12")
    label6.grid(row = 5, column = 0, padx = 2, pady = 2)

    label7 = Label(frame, text = "Day:").grid(row = 0, column = 1, sticky = "w")
    global day4
    day4 = Entry(frame, width = 10)
    day4.grid(row = 0, column = 2, padx = 2, pady = 2)

    label8 = Label(frame, text = "Month:").grid(row = 0, column = 3)
    global month4 
    month4 = Entry(frame, width = 10)
    month4.grid(row = 0, column = 4, padx = 2, pady = 2)

    label9 = Label(frame, text = "Year:").grid(row = 0, column = 5)
    global year4 
    year4 = Entry(frame, width = 10)
    year4.grid(row = 0, column = 6, padx = 2, pady = 2)

    label10 = Label(top4, text = "Source URL:", font="calibri 12")
    label10.grid(row = 6)
    global url4 
    url4 = Entry(top4, width = 100)
    url4.grid(row = 6, column = 1, padx = 2, pady = 2)

    generate = Button(top4, text = "Generate Works Cited Document", font="calibri 12", command = clickgen4)
    generate.grid(row = 7, column = 1, pady = 5)
    proof = Label(top4, text = "Please proofread your document...iFormat is not responsible for your bad grades", font="calibri 10 italic")
    proof.grid(row = 8, column = 1)

def click5():
    global top5 # Online Video
    top5 = Toplevel()
    top5.title("Online Video...Capitalization and formatting will be automatically applied")

    label1 = Label(top5, text = "Author's First Name:", font="calibri 12")
    label1.grid(row = 0)
    global authorf5
    authorf5 = Entry(top5, width = 100)
    authorf5.grid(row = 0, column = 1, padx = 2, pady = 2)

    label2 = Label(top5, text = "Author's Last Name:", font="calibri 12")
    label2.grid(row = 1)
    global authorl5
    authorl5 = Entry(top5, width = 100)
    authorl5.grid(row = 1, column = 1, padx = 2, pady = 2)

    label3 = Label(top5, text = "Video Title:", font="calibri 12")
    label3.grid(row = 2)
    global vid5
    vid5 = Entry(top5, width = 100)
    vid5.grid(row = 2, column = 1, padx = 2, pady = 2)

    label4 = Label(top5, text = "Host Website:", font="calibri 12")
    label4.grid(row = 3)
    global website5
    website5 = Entry(top5, width = 100)
    website5.grid(row = 3, column = 1, padx = 2, pady = 2)

    frame = Frame(top5)
    frame.grid(row = 4, column = 1, columnspan = 6, sticky = "w")
    label4_1 = Label(top5, text = "Date Published:", font="calibri 12")
    label4_1.grid(row = 4, column = 0, padx = 2, pady = 2)

    label5 = Label(frame, text = "Day:").grid(row = 0, column = 1, sticky = "w")
    global day5
    day5 = Entry(frame, width = 10)
    day5.grid(row = 0, column = 2, padx = 2, pady = 2)

    label6 = Label(frame, text = "Month:").grid(row = 0, column = 3)
    global month5 
    month5 = Entry(frame, width = 10)
    month5.grid(row = 0, column = 4, padx = 2, pady = 2)

    label7 = Label(frame, text = "Year:").grid(row = 0, column = 5)
    global year5 
    year5 = Entry(frame, width = 10)
    year5.grid(row = 0, column = 6, padx = 2, pady = 2)

    label8 = Label(top5, text = "Video URL:", font="calibri 12")
    label8.grid(row = 5)
    global url5 
    url5 = Entry(top5, width = 100)
    url5.grid(row = 5, column = 1, padx = 2, pady = 2)

    generate = Button(top5, text = "Generate Works Cited Document", font="calibri 12", command = clickgen5)
    generate.grid(row = 6, column = 1, pady = 5)
    proof = Label(top5, text = "Please proofread your document...iFormat is not responsible for your bad grades", font="calibri 10 italic")
    proof.grid(row = 7, column = 1)

def click6(): # Online Database
    global top6 
    top6 = Toplevel()
    top6.title("Online Database...Capitalization and formatting will be automatically applied")

    label1 = Label(top6, text = "Author's First Name:", font="calibri 12")
    label1.grid(row = 0)
    global authorf6
    authorf6 = Entry(top6, width = 100)
    authorf6.grid(row = 0, column = 1, padx = 2, pady = 2)

    label2 = Label(top6, text = "Author's Last Name:", font="calibri 12")
    label2.grid(row = 1)
    global authorl6
    authorl6 = Entry(top6, width = 100)
    authorl6.grid(row = 1, column = 1, padx = 2, pady = 2)

    label3 = Label(top6, text = "Work/Article Title:", font="calibri 12")
    label3.grid(row = 2)
    global work6
    work6 = Entry(top6, width = 100)
    work6.grid(row = 2, column = 1, padx = 2, pady = 2)

    label4 = Label(top6, text = "Title of Periodical:", font="calibri 12")
    label4.grid(row = 3)
    global per_title6 
    per_title6 = Entry(top6, width = 100)
    per_title6.grid(row = 3, column = 1, padx = 2, pady = 2)

    frame = Frame(top6)
    frame.grid(row = 4, column = 1, columnspan = 6, sticky = "w")
    label4_1 = Label(top6, text = "Volume/Issue and Year:", font="calibri 12")
    label4_1.grid(row = 4, column = 0, padx = 2, pady = 2)

    label5 = Label(frame, text = "Volume:").grid(row = 0, column = 1, sticky = "w")
    global volume6
    volume6 = Entry(frame, width = 10)
    volume6.grid(row = 0, column = 2, padx = 2, pady = 2)

    label6 = Label(frame, text = "Issue:").grid(row = 0, column = 3)
    global issue6 
    issue6 = Entry(frame, width = 10)
    issue6.grid(row = 0, column = 4, padx = 2, pady = 2)

    label7 = Label(frame, text = "Year:").grid(row = 0, column = 5)
    global year6 
    year6 = Entry(frame, width = 10)
    year6.grid(row = 0, column = 6, padx = 2, pady = 2)

    frame2 = Frame(top6)
    frame2.grid(row = 5, column = 1, columnspan = 4, sticky = "w")
    label7_1 = Label(top6, text = "Page Numbers:", font="calibri 12")
    label7_1.grid(row = 5, column = 0, padx = 2, pady = 2)

    label8 = Label(frame2, text = "Start Page:").grid(row = 0, column = 1)
    global pgstart6
    pgstart6 = Entry(frame2, width = 10)
    pgstart6.grid(row = 0, column = 2, padx = 2, pady = 2)

    label9 = Label(frame2, text = "End Page:").grid(row = 0, column = 3)
    global pgend6
    pgend6 = Entry(frame2, width = 10)
    pgend6.grid(row = 0, column = 4, padx = 2, pady = 2)

    label10 = Label(top6, text = "Name of Database:", font="calibri 12")
    label10.grid(row = 6)
    global title6
    title6 = Entry(top6, width = 100)
    title6.grid(row = 6, column = 1, padx = 2, pady = 2)

    label11 = Label(top6, text = "DOI or Permalink:", font="calibri 12")
    label11.grid(row = 7)
    global doi6 
    doi6 = Entry(top6, width = 100)
    doi6.grid(row = 7, column = 1, padx = 2, pady = 2)

    generate = Button(top6, text = "Generate Works Cited Document", font="calibri 12", command = clickgen6)
    generate.grid(row = 8, column = 1, pady = 5)
    proof = Label(top6, text = "Please proofread your document...iFormat is not responsible for your bad grades", font="calibri 10 italic")
    proof.grid(row = 9, column = 1)

def click7(): # Film or Movie
    global top7
    top7 = Toplevel()
    top7.title("Film or Movie...Capitalization and formatting will be automatically applied")

    label1 = Label(top7, text = "Actor/Actress First Name:", font="calibri 12")
    label1.grid(row = 0)
    global perff7
    perff7 = Entry(top7, width = 100)
    perff7.grid(row = 0, column = 1, padx = 2, pady = 2)

    label2 = Label(top7, text = "Actor/Actress Last Name:", font="calibri 12")
    label2.grid(row = 1)
    global perfl7
    perfl7 = Entry(top7, width = 100)
    perfl7.grid(row = 1, column = 1, padx = 2, pady = 2)

    label3 = Label(top7, text = "Film/Movie Title:", font="calibri 12")
    label3.grid(row = 2)
    global title7
    title7 = Entry(top7, width = 100)
    title7.grid(row = 2, column = 1, padx = 2, pady = 2)

    label4 = Label(top7, text = "Director Full Name:", font="calibri 12")
    label4.grid(row = 3)
    global dir7
    dir7 = Entry(top7, width = 100)
    dir7.grid(row = 3, column = 1, padx = 2, pady = 2)

    label5 = Label(top7, text = "Distributor/Studio:", font="calibri 12")
    label5.grid(row = 4)
    global stud7
    stud7 = Entry(top7, width = 100)
    stud7.grid(row = 4, column = 1, padx = 2, pady = 2)

    label6 = Label(top7, text = "Year of Release:", font="calibri 12")
    label6.grid(row = 5)
    global year7
    year7 = Entry(top7, width = 100)
    year7.grid(row = 5, column = 1, padx = 2, pady = 2)

    generate = Button(top7, text = "Generate Works Cited Document", font="calibri 12", command = clickgen7)
    generate.grid(row = 7, column = 1, pady = 5)
    proof = Label(top7, text = "Please proofread your document...iFormat is not responsible for your bad grades", font="calibri 10 italic")
    proof.grid(row = 8, column = 1)

def click8(): # Wikipedia
    global top8 
    top8 = Toplevel()
    top8.title("Wikipedia...Capitalization and formatting will be automatically applied")

    label1 = Label(top8, text = "Entry/Subject:", font="calibri 12")
    label1.grid(row = 0)
    global sub8
    sub8 = Entry(top8, width = 100)
    sub8.grid(row = 0, column = 1, padx = 2, pady = 2)

    frame = Frame(top8)
    frame.grid(row = 1, column = 1, columnspan = 6, sticky = "w")
    label1_1 = Label(top8, text = "Date Last Updated:", font="calibri 12")
    label1_1.grid(row = 1, column = 0, padx = 2, pady = 2)

    label2 = Label(frame, text = "Day:").grid(row = 0, column = 1, sticky = "w")
    global day8
    day8 = Entry(frame, width = 10)
    day8.grid(row = 0, column = 2, padx = 2, pady = 2)

    label3 = Label(frame, text = "Month:").grid(row = 0, column = 3)
    global month8 
    month8 = Entry(frame, width = 10)
    month8.grid(row = 0, column = 4, padx = 2, pady = 2)

    label4 = Label(frame, text = "Year:").grid(row = 0, column = 5)
    global year8 
    year8 = Entry(frame, width = 10)
    year8.grid(row = 0, column = 6, padx = 2, pady = 2)

    label5 = Label(top8, text = "Permalink/URL:", font="calibri 12")
    label5.grid(row = 2)
    global url8
    url8 = Entry(top8, width = 100)
    url8.grid(row= 2, column = 1, padx = 2, pady = 2)

    generate = Button(top8, text = "Generate Works Cited Document", font="calibri 12", command = clickgen8)
    generate.grid(row = 3, column = 1, pady = 5)
    proof = Label(top8, text = "Please proofread your document...iFormat is not responsible for your bad grades", font="calibri 10 italic")
    proof.grid(row = 4, column = 1)
    
heading = Label(root, text = "iFormat MLA Generator", font="calibri 20 bold")
subhead = Label(root, text = "Please select your source material\n", font="calibri 14 italic") 
heading.grid(row = 0, columnspan = 4) 
subhead.grid(row = 1, columnspan = 4)

source_type1 = Button(root, text = "Print Book", font="calibri 12", command = click1)
source_type2 = Button(root, text = "Print Article", font="calibri 12", command = click2)
source_type3 = Button(root, text = "Website", font="calibri 12", command = click3)
source_type4 = Button(root, text = "Web Article", font="calibri 12", command = click4)
source_type5 = Button(root, text = "Online Video", font="calibri 12", command = click5)
source_type6 = Button(root, text = "Online Database", font="calibri 12", command = click6)
source_type7 = Button(root, text = "Film or Movie", font="calibri 12", command = click7)
source_type8 = Button(root, text = "Wikipedia", font="calibri 12", command = click8)

source_type1.config(height = 1, width = 15)
source_type2.config(height = 1, width = 15)
source_type3.config(height = 1, width = 15)
source_type4.config(height = 1, width = 15)
source_type5.config(height = 1, width = 15)
source_type6.config(height = 1, width = 15)
source_type7.config(height = 1, width = 15)
source_type8.config(height = 1, width = 15)

source_type1.grid(row = 4, column = 0)
source_type2.grid(row = 4, column = 1)
source_type3.grid(row = 4, column = 2)
source_type4.grid(row = 4, column = 3)
source_type5.grid(row = 5, column = 0)
source_type6.grid(row = 5, column = 1)
source_type7.grid(row = 5, column = 2)
source_type8.grid(row = 5, column = 3)

root.mainloop() 
